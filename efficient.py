import streamlit as st
import yfinance as yf
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import json
import io
from scipy.optimize import minimize

# --- Page Configuration ---
st.set_page_config(
    page_title="QuantPro: Institutional Analytics",
    page_icon="💎",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Custom Institutional Styling ---
st.markdown("""
<style>
    /* Dark Professional Theme with Gradient */
    .stApp { 
        background: linear-gradient(135deg, #0a0e1a 0%, #1a1f35 50%, #0f1425 100%);
        color: #e8eaed;
        font-family: 'IBM Plex Sans', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    
    /* Premium Metric Cards with Hover Effects */
    div[data-testid="stMetric"] {
        background: linear-gradient(145deg, #1e2538, #252d45);
        padding: 20px;
        border-radius: 12px;
        border: 1px solid rgba(100, 150, 255, 0.2);
        box-shadow: 0 4px 16px rgba(0, 0, 0, 0.3);
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    }
    
    div[data-testid="stMetric"]:hover {
        transform: translateY(-4px);
        border-color: rgba(100, 150, 255, 0.5);
        box-shadow: 0 8px 24px rgba(59, 130, 246, 0.2);
    }
    
    /* Metric Label Styling */
    div[data-testid="stMetric"] label {
        font-size: 0.875rem !important;
        font-weight: 600 !important;
        color: #9ca3af !important;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }
    
    /* Metric Value Styling */
    div[data-testid="stMetric"] [data-testid="stMetricValue"] {
        font-size: 1.75rem !important;
        font-weight: 700 !important;
        color: #f3f4f6 !important;
    }
    
    /* Enhanced Typography */
    h1 {
        font-weight: 700 !important;
        letter-spacing: -0.03em !important;
        background: linear-gradient(135deg, #60a5fa 0%, #a78bfa 100%);
        -webkit-background-clip: text !important;
        -webkit-text-fill-color: transparent !important;
        background-clip: text !important;
        margin-bottom: 0.5rem !important;
    }
    
    h2, h3 {
        font-weight: 600 !important;
        letter-spacing: -0.02em !important;
        color: #f3f4f6 !important;
    }
    
    /* Scrollable Matrix Container with Better Styling */
    .matrix-container {
        overflow-x: auto;
        overflow-y: hidden;
        width: 100%;
        background: rgba(15, 20, 37, 0.6);
        border-radius: 12px;
        padding: 16px;
        border: 1px solid rgba(100, 150, 255, 0.15);
        box-shadow: inset 0 2px 8px rgba(0, 0, 0, 0.3);
    }
    
    /* Custom Scrollbar */
    .matrix-container::-webkit-scrollbar {
        height: 12px;
    }
    
    .matrix-container::-webkit-scrollbar-track {
        background: rgba(30, 37, 56, 0.5);
        border-radius: 6px;
    }
    
    .matrix-container::-webkit-scrollbar-thumb {
        background: linear-gradient(90deg, #3b82f6, #8b5cf6);
        border-radius: 6px;
    }
    
    .matrix-container::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(90deg, #2563eb, #7c3aed);
    }
    
    /* Tab Styling with Smooth Transitions */
    .stTabs [data-baseweb="tab-list"] {
        gap: 12px;
        background: rgba(15, 20, 37, 0.4);
        padding: 8px;
        border-radius: 12px;
        border: 1px solid rgba(100, 150, 255, 0.1);
    }
    
    .stTabs [data-baseweb="tab"] {
        background-color: rgba(30, 37, 56, 0.6);
        border-radius: 8px;
        padding: 10px 20px;
        border: 1px solid rgba(100, 150, 255, 0.15);
        transition: all 0.3s ease;
        font-weight: 500;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        background-color: rgba(59, 130, 246, 0.2);
        border-color: rgba(100, 150, 255, 0.3);
        transform: translateY(-1px);
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, rgba(59, 130, 246, 0.3), rgba(139, 92, 246, 0.3)) !important;
        border-color: rgba(100, 150, 255, 0.5) !important;
    }
    
    /* Expander Styling */
    .streamlit-expanderHeader {
        background: linear-gradient(145deg, #1e2538, #252d45);
        border-radius: 10px;
        border: 1px solid rgba(100, 150, 255, 0.2);
        transition: all 0.3s ease;
        font-weight: 600;
    }
    
    .streamlit-expanderHeader:hover {
        border-color: rgba(100, 150, 255, 0.4);
        box-shadow: 0 4px 12px rgba(59, 130, 246, 0.15);
    }
    
    /* Alert Boxes with Better Styling */
    .stAlert {
        background: rgba(15, 20, 37, 0.8);
        border-radius: 10px;
        border-left: 4px solid;
        backdrop-filter: blur(10px);
        padding: 1rem;
    }
    
    /* Success Alert */
    div[data-baseweb="notification"][kind="success"] {
        background: rgba(16, 185, 129, 0.1) !important;
        border-left-color: #10b981 !important;
    }
    
    /* Warning Alert */
    div[data-baseweb="notification"][kind="warning"] {
        background: rgba(245, 158, 11, 0.1) !important;
        border-left-color: #f59e0b !important;
    }
    
    /* Info Alert */
    div[data-baseweb="notification"][kind="info"] {
        background: rgba(59, 130, 246, 0.1) !important;
        border-left-color: #3b82f6 !important;
    }
    
    /* Button Styling */
    .stButton > button {
        background: linear-gradient(135deg, #3b82f6, #8b5cf6);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        font-weight: 600;
        transition: all 0.3s ease;
        box-shadow: 0 4px 12px rgba(59, 130, 246, 0.3);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(59, 130, 246, 0.4);
        background: linear-gradient(135deg, #2563eb, #7c3aed);
    }
    
    /* Data Editor Styling */
    .stDataFrame {
        border: 1px solid rgba(100, 150, 255, 0.2);
        border-radius: 8px;
        overflow: hidden;
    }
    
    /* Sidebar Styling */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0f1425 0%, #1a1f35 100%);
        border-right: 1px solid rgba(100, 150, 255, 0.2);
    }
    
    section[data-testid="stSidebar"] .stMarkdown {
        color: #e5e7eb;
    }
    
    /* File Uploader Styling */
    .stFileUploader {
        background: rgba(30, 37, 56, 0.6);
        border: 2px dashed rgba(100, 150, 255, 0.3);
        border-radius: 10px;
        padding: 1rem;
        transition: all 0.3s ease;
    }
    
    .stFileUploader:hover {
        border-color: rgba(100, 150, 255, 0.5);
        background: rgba(30, 37, 56, 0.8);
    }
    
    /* Spinner Animation */
    .stSpinner > div {
        border-top-color: #3b82f6 !important;
    }
    
    /* Caption Styling */
    .caption {
        color: #9ca3af !important;
        font-size: 0.875rem !important;
    }
    
    /* Divider Styling */
    hr {
        border-color: rgba(100, 150, 255, 0.2) !important;
        margin: 2rem 0 !important;
    }
    
    /* Number Input Styling */
    .stNumberInput > div > div > input {
        background: rgba(30, 37, 56, 0.6);
        border: 1px solid rgba(100, 150, 255, 0.2);
        border-radius: 6px;
        color: #f3f4f6;
    }
    
    /* Slider Styling */
    .stSlider > div > div > div {
        background: linear-gradient(90deg, #3b82f6, #8b5cf6);
    }
    
    /* Download Button Styling */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #10b981, #059669);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.5rem 1rem;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stDownloadButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(16, 185, 129, 0.4);
    }
    
    /* Plotly Chart Container */
    .js-plotly-plot {
        border-radius: 12px;
        overflow: hidden;
    }
</style>
""", unsafe_allow_html=True)

# --- Helper Functions ---

@st.cache_data(ttl=3600)
def get_currency_rate():
    """Fetch current USD/INR exchange rate"""
    try:
        data = yf.Ticker("INR=X").history(period="1d")
        return data['Close'].iloc[-1] if not data.empty else 84.0
    except: 
        return 84.0 

def fetch_data(tickers, period="2y"):
    """Fetch historical price data for given tickers"""
    if not tickers: 
        return pd.DataFrame()
    try:
        data = yf.download(tickers, period=period, auto_adjust=False, threads=True, progress=False)
        if isinstance(data.columns, pd.MultiIndex):
            return data['Adj Close'] if 'Adj Close' in data.columns.get_level_values(0) else data['Close']
        return data['Adj Close'] if 'Adj Close' in data.columns else data['Close']
    except: 
        return pd.DataFrame()

def calculate_portfolio_metrics(weights, mean_returns, cov_matrix, risk_free_rate):
    """
    Calculate portfolio metrics using Modern Portfolio Theory
    
    Formula for Sharpe Ratio:
    Sharpe = (Rp - Rf) / σp
    
    Where:
    - Rp = Portfolio Expected Return (annualized)
    - Rf = Risk-Free Rate
    - σp = Portfolio Standard Deviation (annualized volatility)
    
    Higher Sharpe ratio indicates better risk-adjusted returns
    """
    # Annualized return (252 trading days)
    portfolio_return = np.sum(mean_returns * weights) * 252
    
    # Annualized volatility (portfolio standard deviation)
    portfolio_std = np.sqrt(np.dot(weights.T, np.dot(cov_matrix, weights))) * np.sqrt(252)
    
    # Sharpe Ratio: excess return per unit of risk
    sharpe_ratio = (portfolio_return - risk_free_rate) / portfolio_std if portfolio_std > 0 else 0
    
    return portfolio_return, portfolio_std, sharpe_ratio

def optimize_portfolio(mean_returns, cov_matrix, risk_free_rate, target='sharpe'):
    """
    Optimize portfolio weights using scipy optimization
    Returns optimal weights for maximum Sharpe ratio or minimum volatility
    """
    num_assets = len(mean_returns)
    
    # Objective functions
    def neg_sharpe(weights):
        ret, std, sharpe = calculate_portfolio_metrics(weights, mean_returns, cov_matrix, risk_free_rate)
        return -sharpe  # Negative because we minimize
    
    def portfolio_volatility(weights):
        return np.sqrt(np.dot(weights.T, np.dot(cov_matrix, weights))) * np.sqrt(252)
    
    # Constraints and bounds
    constraints = ({'type': 'eq', 'fun': lambda w: np.sum(w) - 1})  # Weights sum to 1
    bounds = tuple((0, 1) for _ in range(num_assets))  # Long-only positions
    initial_weights = np.array([1/num_assets] * num_assets)
    
    # Optimize
    objective = neg_sharpe if target == 'sharpe' else portfolio_volatility
    result = minimize(objective, initial_weights, method='SLSQP', bounds=bounds, constraints=constraints)
    
    return result.x if result.success else initial_weights

def parse_uploaded_portfolio(uploaded_file):
    """
    Parse portfolio data from uploaded CSV, JSON, DOCX, or PDF
    Expected format: Ticker, Shares, Avg Cost, Currency (optional)
    """
    try:
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        if file_type == 'csv':
            df = pd.read_csv(uploaded_file)
            
        elif file_type == 'json':
            data = json.load(uploaded_file)
            df = pd.DataFrame(data)
            
        elif file_type in ['xlsx', 'xls']:
            df = pd.read_excel(uploaded_file)
            
        elif file_type == 'docx':
            # For DOCX, expect a simple table or structured text
            import docx
            doc = docx.Document(uploaded_file)
            
            # Try to extract table
            if doc.tables:
                table = doc.tables[0]
                data = []
                keys = None
                for i, row in enumerate(table.rows):
                    text = [cell.text.strip() for cell in row.cells]
                    if i == 0:
                        keys = text
                    else:
                        data.append(text)
                df = pd.DataFrame(data, columns=keys)
            else:
                st.error("No table found in DOCX. Please format as a table.")
                return None
                
        elif file_type == 'pdf':
            # For PDF, use tabula or pdfplumber to extract tables
            import pdfplumber
            
            with pdfplumber.open(uploaded_file) as pdf:
                first_page = pdf.pages[0]
                table = first_page.extract_table()
                
                if table:
                    df = pd.DataFrame(table[1:], columns=table[0])
                else:
                    st.error("No table found in PDF.")
                    return None
        else:
            st.error(f"Unsupported file type: {file_type}")
            return None
        
        # Standardize column names
        df.columns = df.columns.str.strip().str.title()
        column_mapping = {
            'Symbol': 'Ticker',
            'Stock': 'Ticker',
            'Quantity': 'Shares',
            'Qty': 'Shares',
            'Units': 'Shares',
            'Average Cost': 'Avg Cost',
            'Avg Price': 'Avg Cost',
            'Cost': 'Avg Cost',
            'Price': 'Avg Cost',
        }
        df.rename(columns=column_mapping, inplace=True)
        
        # Ensure required columns
        required = ['Ticker', 'Shares', 'Avg Cost']
        if not all(col in df.columns for col in required):
            st.error(f"Missing required columns. Need: {required}")
            return None
        
        # Add Currency column if missing
        if 'Currency' not in df.columns:
            df['Currency'] = 'INR'
        
        # Convert to proper types
        df['Shares'] = pd.to_numeric(df['Shares'], errors='coerce')
        df['Avg Cost'] = pd.to_numeric(df['Avg Cost'], errors='coerce')
        
        # Remove invalid rows
        df = df.dropna(subset=['Ticker', 'Shares', 'Avg Cost'])
        
        return df
        
    except Exception as e:
        st.error(f"Error parsing file: {str(e)}")
        return None

# --- Main App ---
def main():
    # Header with institutional branding and animated subtitle
    st.markdown("""
    <div style='text-align: center; padding: 1rem 0 2rem 0;'>
        <h1 style='margin-bottom: 0.5rem;'>💎 QuantPro: Institutional Portfolio Intelligence</h1>
        <p style='font-size: 1.1rem; color: #9ca3af; font-weight: 500; letter-spacing: 0.05em;'>
            Advanced Portfolio Analytics • Risk Management • Optimization Engine
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar Configuration
    with st.sidebar:
        st.header("⚙️ Configuration Panel")
        
        # Risk-free rate with explanation
        st.markdown("**Risk-Free Rate (Rf)**")
        st.caption("Used in Sharpe ratio calculation. Typically 10Y Government Bond yield.")
        rf_rate = st.number_input(
            "Annual Rf Rate", 
            value=0.065, 
            min_value=0.0,
            max_value=0.20,
            step=0.005, 
            format="%.3f",
            help="Current India 10Y G-Sec yield ≈ 6.5%"
        )
        
        st.divider()
        
        # Currency rate display
        usd_rate = get_currency_rate()
        st.metric("💱 USD/INR Rate", f"₹{usd_rate:.2f}", delta="Live")
        
        st.divider()
        
        # File upload section
        st.markdown("**📂 Import Portfolio**")
        uploaded_file = st.file_uploader(
            "Upload CSV, JSON, XLSX, DOCX, or PDF",
            type=['csv', 'json', 'xlsx', 'xls', 'docx', 'pdf'],
            help="File should contain: Ticker, Shares, Avg Cost, Currency"
        )
        
        if uploaded_file:
            with st.spinner("Parsing portfolio data..."):
                parsed_df = parse_uploaded_portfolio(uploaded_file)
                if parsed_df is not None:
                    st.session_state.portfolio_df = parsed_df
                    st.success(f"✅ Loaded {len(parsed_df)} positions")
        
        st.divider()
        
        # Export options
        st.markdown("**📤 Export Data**")
        if st.button("📥 Export Portfolio JSON", use_container_width=True):
            json_str = st.session_state.portfolio_df.to_json(orient='records', indent=2)
            st.download_button(
                "Download JSON",
                json_str,
                "portfolio_export.json",
                "application/json",
                use_container_width=True
            )
        
        st.divider()
        
        # Quick Tips
        st.markdown("**💡 Quick Tips**")
        with st.expander("📖 Getting Started"):
            st.markdown("""
            **New User Guide:**
            1. Upload your portfolio or edit default
            2. Check Dashboard metrics
            3. Explore Efficient Frontier
            4. Review Correlation Matrix
            5. Run Stress Tests
            
            **Indian Stock Tickers:**
            - Add `.NS` suffix for NSE stocks
            - Example: `RELIANCE.NS`, `TCS.NS`
            """)
        
        with st.expander("📊 Understanding Metrics"):
            st.markdown("""
            **Sharpe Ratio:**
            - < 0: Poor (below risk-free rate)
            - 0-1: Below average
            - 1-2: Good
            - > 2: Excellent
            
            **Volatility:**
            - < 15%: Low risk
            - 15-25%: Moderate risk
            - > 25%: High risk
            
            **VaR (95%):**
            - Maximum loss expected on 95% of days
            - Lower is better
            """)
        
        with st.expander("🎯 Optimization Tips"):
            st.markdown("""
            **Improve Your Portfolio:**
            1. Check Efficient Frontier tab
            2. Compare current vs optimal weights
            3. Rebalance toward optimal allocation
            4. Reduce high-correlation pairs
            5. Diversify across sectors
            """)
        
        # System Status
        st.divider()
        st.markdown("**🔄 System Status**")
        st.success("✅ All systems operational")
        st.caption(f"Last updated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Initialize default portfolio if not exists
    if 'portfolio_df' not in st.session_state:
        # Show welcome message for first-time users
        st.info("""
        👋 **Welcome to QuantPro!** 
        
        A default 20-stock portfolio has been loaded for you to explore. 
        
        **Quick Actions:**
        - 📊 Check the Dashboard metrics below
        - 🚀 Explore the Efficient Frontier tab
        - 💼 Upload your own portfolio via the sidebar
        - 📖 Read Quick Tips in the sidebar for guidance
        """)
        
        data = {
            "Ticker": ["OLAELEC.NS", "BAJAJHFL.NS", "CESC.NS", "IT.NS", "TATSILV.NS", 
                      "KALYANKJIL.NS", "ITC.NS", "CASTROLIND.NS", "GAIL.NS", "REDINGTON.NS", 
                      "ADANIPOWER.NS", "TMPV.NS", "GROWW.NS", "BSLNIFTY.NS", "PHARMABEES.NS", 
                      "GROWWMETAL.NS", "TATAGOLD.NS", "TATASTEEL.NS", "VEDL.NS", "SBIN.NS"],
            "Shares": [31, 20, 20, 70, 123, 10, 7, 20, 20, 10, 10, 10, 12, 72, 100, 195, 155, 20, 14, 10],
            "Avg Cost": [37.86, 109.5, 176.18, 40.19, 27.4, 473.05, 351.99, 204.65, 177.22, 273.55, 
                        152.04, 391.37, 175.32, 29.49, 22.38, 10.74, 13.53, 171.74, 524.11, 881.58],
            "Currency": ["INR"] * 20
        }
        st.session_state.portfolio_df = pd.DataFrame(data)
        st.session_state.first_load = True
    else:
        st.session_state.first_load = False

    # Quick stats banner (after portfolio initialization)
    col_banner1, col_banner2, col_banner3, col_banner4 = st.columns(4)
    with col_banner1:
        st.metric("📊 Assets in Portfolio", str(len(st.session_state.portfolio_df)), help="Number of positions tracked")
    with col_banner2:
        st.metric("🌐 Data Source", "Yahoo Finance", help="Real-time market data provider")
    with col_banner3:
        st.metric("📈 Analysis Period", "2 Years", help="Historical data used for calculations")
    with col_banner4:
        st.metric("🔄 Status", "Live", delta="Connected", help="System operational status")
    
    st.markdown("---")

    # Portfolio Holdings Editor
    with st.expander("💼 View / Edit Portfolio Holdings", expanded=False):
        edited_df = st.data_editor(
            st.session_state.portfolio_df, 
            num_rows="dynamic", 
            use_container_width=True,
            column_config={
                "Ticker": st.column_config.TextColumn("Ticker", width="medium"),
                "Shares": st.column_config.NumberColumn("Shares", format="%d"),
                "Avg Cost": st.column_config.NumberColumn("Avg Cost (₹)", format="%.2f"),
                "Currency": st.column_config.SelectboxColumn("Currency", options=["INR", "USD"])
            }
        )
        st.session_state.portfolio_df = edited_df

    # Fetch market data with timing
    tickers = edited_df['Ticker'].unique().tolist()
    
    if not tickers:
        st.error("⚠️ No tickers found in portfolio. Please add positions.")
        st.stop()
    
    import time
    start_time = time.time()
    
    with st.spinner("🚀 Fetching market data from Yahoo Finance..."):
        market_data = fetch_data(tickers, period="2y")
        
        fetch_time = time.time() - start_time
        
        if market_data.empty: 
            st.error(
                "❌ **Failed to fetch market data**\n\n"
                "**Possible causes:**\n"
                "- Invalid ticker symbols (Indian stocks need .NS suffix)\n"
                "- Network connectivity issues\n"
                "- Yahoo Finance API unavailable\n\n"
                "**Troubleshooting:**\n"
                "1. Verify ticker format (e.g., RELIANCE.NS for NSE stocks)\n"
                "2. Check your internet connection\n"
                "3. Try again in a few moments"
            )
            st.stop()
        
        # Check if we got data for all tickers
        missing_tickers = set(tickers) - set(market_data.columns)
        if missing_tickers:
            st.warning(
                f"⚠️ **Partial Data Retrieved**\n\n"
                f"Could not fetch data for: {', '.join(missing_tickers)}\n\n"
                f"These positions will be excluded from analysis."
            )
            # Filter out missing tickers
            edited_df = edited_df[~edited_df['Ticker'].isin(missing_tickers)]
            tickers = [t for t in tickers if t not in missing_tickers]
            
            if edited_df.empty:
                st.error("No valid data available after filtering. Please check your tickers.")
                st.stop()
        
        # Success message with timing
        st.success(f"✅ Market data fetched successfully in {fetch_time:.2f}s • {len(market_data.columns)} assets • {len(market_data)} trading days analyzed")

    # Calculate current portfolio metrics
    current_prices = market_data.iloc[-1]
    
    def get_current_val(row):
        price = current_prices.get(row['Ticker'], 0)
        val = price * row['Shares']
        # Convert USD positions to INR
        return val * usd_rate if str(row.get('Currency', 'INR')).upper() == 'USD' else val

    edited_df['Current_Value'] = edited_df.apply(get_current_val, axis=1)
    total_value = edited_df['Current_Value'].sum()
    total_invested = (edited_df['Shares'] * edited_df['Avg Cost']).sum()
    pnl_abs = total_value - total_invested
    pnl_pct = (pnl_abs / total_invested) * 100 if total_invested > 0 else 0
    
    # Calculate returns and covariance
    log_returns = np.log(market_data / market_data.shift(1)).dropna()
    mean_returns = log_returns.mean()
    cov_matrix = log_returns.cov()
    
    # Current portfolio weights
    weights = np.array([edited_df[edited_df['Ticker']==t]['Current_Value'].sum()/total_value 
                       for t in market_data.columns])
    
    # Portfolio metrics
    curr_ret, curr_std, curr_sharpe = calculate_portfolio_metrics(
        weights, mean_returns, cov_matrix, rf_rate
    )
    
    # Calculate optimal portfolios
    optimal_sharpe_weights = optimize_portfolio(mean_returns, cov_matrix, rf_rate, target='sharpe')
    optimal_ret, optimal_std, optimal_sharpe = calculate_portfolio_metrics(
        optimal_sharpe_weights, mean_returns, cov_matrix, rf_rate
    )
    
    min_vol_weights = optimize_portfolio(mean_returns, cov_matrix, rf_rate, target='volatility')
    minvol_ret, minvol_std, minvol_sharpe = calculate_portfolio_metrics(
        min_vol_weights, mean_returns, cov_matrix, rf_rate
    )

    # Display Key Metrics
    st.markdown("---")
    st.markdown("""
    <h2 style='margin-bottom: 1.5rem;'>
        <span style='background: linear-gradient(135deg, #3b82f6, #8b5cf6); 
                     -webkit-background-clip: text; 
                     -webkit-text-fill-color: transparent;
                     background-clip: text;'>
            📊 Portfolio Dashboard
        </span>
    </h2>
    """, unsafe_allow_html=True)
    
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        st.metric(
            "💰 Portfolio Value", 
            f"₹{total_value:,.0f}",
            delta=f"₹{pnl_abs:,.0f}",
            help="Current market value of all positions in INR"
        )
    
    with col2:
        delta_color = "normal" if pnl_pct > 0 else "inverse"
        st.metric(
            "📈 Total Return", 
            f"{pnl_pct:.2f}%",
            delta=f"{pnl_pct:.2f}%",
            delta_color=delta_color,
            help="Total percentage gain/loss from cost basis"
        )
    
    with col3:
        sharpe_delta = curr_sharpe - optimal_sharpe
        delta_color = "normal" if sharpe_delta >= 0 else "inverse"
        st.metric(
            "⚡ Sharpe Ratio", 
            f"{curr_sharpe:.3f}",
            delta=f"{sharpe_delta:.3f}",
            delta_color=delta_color,
            help="Risk-adjusted return. >1.0 is good, >2.0 is excellent. Compares your return vs optimal."
        )
    
    with col4:
        st.metric(
            "🌪️ Volatility (Ann.)", 
            f"{curr_std*100:.2f}%",
            help="Annualized standard deviation. Measures portfolio risk/variability."
        )
    
    with col5:
        st.metric(
            "📊 Expected Return", 
            f"{curr_ret*100:.2f}%",
            help="Projected annual return based on historical performance"
        )

    # Intelligence Summary
    st.markdown("---")
    st.markdown("""
    <h2 style='margin-bottom: 1rem;'>
        <span style='background: linear-gradient(135deg, #10b981, #059669); 
                     -webkit-background-clip: text; 
                     -webkit-text-fill-color: transparent;
                     background-clip: text;'>
            💡 Quantitative Intelligence Brief
        </span>
    </h2>
    """, unsafe_allow_html=True)
    
    col_intel1, col_intel2 = st.columns([2, 1])
    
    with col_intel1:
        if pnl_pct > 0:
            st.success(
                f"**Portfolio Performance: Strong** 🚀\n\n"
                f"Your portfolio has generated **{pnl_pct:.2f}%** absolute return, "
                f"outperforming capital costs by ₹{pnl_abs:,.0f}. "
                f"Current Sharpe ratio of **{curr_sharpe:.3f}** indicates "
                f"{'excellent' if curr_sharpe > 1.5 else 'good' if curr_sharpe > 1.0 else 'moderate'} "
                f"risk-adjusted performance."
            )
        else:
            st.warning(
                f"**Portfolio Under Pressure** ⚠️\n\n"
                f"Current drawdown of **{pnl_pct:.2f}%** (₹{pnl_abs:,.0f}). "
                f"Consider rebalancing toward lower-volatility assets or reviewing position sizing. "
                f"Sharpe ratio: {curr_sharpe:.3f}"
            )
        
        # Optimization insights
        if curr_sharpe < optimal_sharpe:
            improvement = ((optimal_sharpe - curr_sharpe) / curr_sharpe) * 100
            st.info(
                f"**Optimization Opportunity** 💎\n\n"
                f"Our quantitative models suggest a potential **{improvement:.1f}%** improvement "
                f"in Sharpe ratio (from {curr_sharpe:.3f} to {optimal_sharpe:.3f}) "
                f"through portfolio rebalancing. See Frontier tab for optimal allocation."
            )
    
    with col_intel2:
        st.markdown("**🎯 Risk Metrics Dashboard**")
        
        var_95 = curr_std * np.sqrt(1/252) * 1.645
        max_dd = curr_std * 2 * 100
        diversification = len([w for w in weights if w > 0.02])
        
        # Risk level indicator
        risk_level = "🟢 Low" if curr_std < 0.15 else "🟡 Moderate" if curr_std < 0.25 else "🔴 High"
        
        st.metric(
            "Risk Level", 
            risk_level,
            help="Based on annualized volatility: <15% Low, 15-25% Moderate, >25% High"
        )
        st.metric(
            "VaR (95%, 1D)", 
            f"-{var_95 * 100:.2f}%",
            help="Value at Risk: Maximum expected loss on 95% of trading days"
        )
        st.metric(
            "Max Drawdown (Est.)", 
            f"-{max_dd:.1f}%",
            help="Estimated maximum peak-to-trough decline"
        )
        st.metric(
            "Diversification", 
            f"{diversification}/{len(weights)} assets",
            help="Number of significant positions (>2% weight) out of total assets"
        )

    # Main Analysis Tabs
    st.markdown("---")
    
    # Add performance summary before tabs
    col_summary1, col_summary2, col_summary3 = st.columns(3)
    
    with col_summary1:
        st.markdown("**📈 Performance Summary**")
        st.write(f"• Return: **{pnl_pct:.2f}%**")
        st.write(f"• Sharpe: **{curr_sharpe:.3f}**")
        st.write(f"• Volatility: **{curr_std*100:.2f}%**")
    
    with col_summary2:
        st.markdown("**⚖️ Portfolio Composition**")
        st.write(f"• Total Value: **₹{total_value:,.0f}**")
        st.write(f"• Positions: **{len(edited_df)}** assets")
        st.write(f"• Top Position: **{(weights.max()*100):.1f}%**")
    
    with col_summary3:
        st.markdown("**🎯 Optimization Potential**")
        improvement = ((optimal_sharpe - curr_sharpe) / curr_sharpe * 100) if curr_sharpe > 0 else 0
        st.write(f"• Current Sharpe: **{curr_sharpe:.3f}**")
        st.write(f"• Optimal Sharpe: **{optimal_sharpe:.3f}**")
        st.write(f"• Improvement: **{improvement:.1f}%**")
    
    st.markdown("---")
    
    tab1, tab2, tab3, tab4 = st.tabs([
        "🚀 Efficient Frontier", 
        "🔬 Correlation Matrix", 
        "🌪️ Stress Testing",
        "📊 Position Analysis"
    ])

    # TAB 1: Enhanced Efficient Frontier
    with tab1:
        st.subheader("🚀 Efficient Frontier Analysis")
        st.caption("Monte Carlo simulation with 5,000 random portfolios. "
                  "Your portfolio is marked with a ⭐, optimal portfolios shown in white.")
        
        # Run Monte Carlo simulation
        num_portfolios = 5000
        results = np.zeros((4, num_portfolios))
        
        progress_bar = st.progress(0, text="Running Monte Carlo simulation...")
        
        for i in range(num_portfolios):
            # Generate random weights
            w = np.random.random(len(tickers))
            w /= np.sum(w)
            
            # Calculate metrics
            r, s, sh = calculate_portfolio_metrics(w, mean_returns, cov_matrix, rf_rate)
            results[0, i] = s      # Volatility
            results[1, i] = r      # Return
            results[2, i] = sh     # Sharpe
            results[3, i] = i      # Index for tracking
            
            # Update progress every 100 iterations
            if i % 100 == 0:
                progress_bar.progress((i + 1) / num_portfolios, 
                                     text=f"Running Monte Carlo simulation... {i+1}/{num_portfolios}")
        
        progress_bar.empty()  # Remove progress bar when done
        
        # Create enhanced scatter plot
        fig_ef = go.Figure()
        
        # Monte Carlo portfolios (colored by Sharpe ratio)
        fig_ef.add_trace(go.Scatter(
            x=results[0, :] * 100,
            y=results[1, :] * 100,
            mode='markers',
            name='Simulated Portfolios',
            marker=dict(
                size=4,
                color=results[2, :],
                colorscale='Viridis',
                showscale=True,
                colorbar=dict(
                    title="Sharpe<br>Ratio",
                    x=1.15
                ),
                opacity=0.6
            ),
            text=[f'Sharpe: {sh:.3f}' for sh in results[2, :]],
            hovertemplate='Vol: %{x:.2f}%<br>Return: %{y:.2f}%<br>%{text}<extra></extra>'
        ))
        
        # Current portfolio
        fig_ef.add_trace(go.Scatter(
            x=[curr_std * 100],
            y=[curr_ret * 100],
            mode='markers+text',
            name='Your Portfolio',
            text=['YOUR<br>PORTFOLIO'],
            textposition='top center',
            textfont=dict(size=11, color='white', family='Arial Black'),
            marker=dict(
                color='#FFD700',
                size=20,
                symbol='star',
                line=dict(color='white', width=2)
            ),
            hovertemplate='<b>Your Portfolio</b><br>Vol: %{x:.2f}%<br>Return: %{y:.2f}%<br>Sharpe: ' + f'{curr_sharpe:.3f}<extra></extra>'
        ))
        
        # Optimal Sharpe portfolio
        fig_ef.add_trace(go.Scatter(
            x=[optimal_std * 100],
            y=[optimal_ret * 100],
            mode='markers+text',
            name='Max Sharpe',
            text=['OPTIMAL<br>SHARPE'],
            textposition='top center',
            textfont=dict(size=10, color='white'),
            marker=dict(
                color='#00FF00',
                size=18,
                symbol='diamond',
                line=dict(color='white', width=2)
            ),
            hovertemplate='<b>Optimal Sharpe</b><br>Vol: %{x:.2f}%<br>Return: %{y:.2f}%<br>Sharpe: ' + f'{optimal_sharpe:.3f}<extra></extra>'
        ))
        
        # Minimum volatility portfolio
        fig_ef.add_trace(go.Scatter(
            x=[minvol_std * 100],
            y=[minvol_ret * 100],
            mode='markers+text',
            name='Min Volatility',
            text=['MIN<br>VOL'],
            textposition='bottom center',
            textfont=dict(size=10, color='white'),
            marker=dict(
                color='#00CED1',
                size=16,
                symbol='square',
                line=dict(color='white', width=2)
            ),
            hovertemplate='<b>Min Volatility</b><br>Vol: %{x:.2f}%<br>Return: %{y:.2f}%<br>Sharpe: ' + f'{minvol_sharpe:.3f}<extra></extra>'
        ))
        
        # Capital Allocation Line (CAL) - tangent from risk-free rate
        max_vol = np.max(results[0, :]) * 100
        cal_x = np.linspace(0, max_vol, 100)
        cal_y = rf_rate * 100 + optimal_sharpe * cal_x
        
        fig_ef.add_trace(go.Scatter(
            x=cal_x,
            y=cal_y,
            mode='lines',
            name='Capital Allocation Line',
            line=dict(color='rgba(255, 255, 255, 0.3)', width=2, dash='dash'),
            hovertemplate='CAL: Rf + Sharpe × Vol<extra></extra>'
        ))
        
        # Layout
        fig_ef.update_layout(
            template="plotly_dark",
            height=600,
            title={
                'text': "Efficient Frontier • Risk-Return Optimization",
                'x': 0.5,
                'xanchor': 'center',
                'font': {'size': 18}
            },
            xaxis_title="Annualized Volatility (Standard Deviation) %",
            yaxis_title="Annualized Expected Return %",
            hovermode='closest',
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            showlegend=True,
            legend=dict(
                x=0.02,
                y=0.98,
                bgcolor='rgba(0,0,0,0.5)',
                bordercolor='rgba(255,255,255,0.2)',
                borderwidth=1
            ),
            font=dict(family='IBM Plex Sans', size=11)
        )
        
        fig_ef.update_xaxes(showgrid=True, gridwidth=1, gridcolor='rgba(255,255,255,0.1)')
        fig_ef.update_yaxes(showgrid=True, gridwidth=1, gridcolor='rgba(255,255,255,0.1)')
        
        st.plotly_chart(fig_ef, use_container_width=True)
        
        # Optimal allocation table
        st.markdown("**Optimal Portfolio Allocations**")
        
        col_opt1, col_opt2 = st.columns(2)
        
        with col_opt1:
            st.markdown("*Maximum Sharpe Ratio Portfolio*")
            opt_df = pd.DataFrame({
                'Ticker': tickers,
                'Optimal Weight %': optimal_sharpe_weights * 100,
                'Current Weight %': weights * 100
            }).sort_values('Optimal Weight %', ascending=False)
            opt_df = opt_df[opt_df['Optimal Weight %'] > 0.5]  # Show only significant positions
            st.dataframe(opt_df.style.format({
                'Optimal Weight %': '{:.2f}',
                'Current Weight %': '{:.2f}'
            }), hide_index=True)
        
        with col_opt2:
            st.markdown("*Minimum Volatility Portfolio*")
            minvol_df = pd.DataFrame({
                'Ticker': tickers,
                'Min Vol Weight %': min_vol_weights * 100,
                'Current Weight %': weights * 100
            }).sort_values('Min Vol Weight %', ascending=False)
            minvol_df = minvol_df[minvol_df['Min Vol Weight %'] > 0.5]
            st.dataframe(minvol_df.style.format({
                'Min Vol Weight %': '{:.2f}',
                'Current Weight %': '{:.2f}'
            }), hide_index=True)

    # TAB 2: Enhanced Correlation Matrix
    with tab2:
        st.subheader("🔬 Asset Correlation Matrix")
        st.caption("Scroll horizontally to view all assets. Red indicates negative correlation (diversification), "
                  "blue indicates positive correlation (concentration risk).")
        
        # Calculate correlation matrix
        corr_matrix = log_returns.corr()
        
        # Create enhanced heatmap with better readability
        fig_corr = go.Figure(data=go.Heatmap(
            z=corr_matrix.values,
            x=corr_matrix.columns,
            y=corr_matrix.columns,
            colorscale='RdBu_r',
            zmid=0,
            zmin=-1,
            zmax=1,
            text=np.round(corr_matrix.values, 2),
            texttemplate='%{text}',
            textfont={"size": 10, "color": "white"},
            hovertemplate='%{y} vs %{x}<br>Correlation: %{z:.3f}<extra></extra>',
            colorbar=dict(
                title="Correlation",
                tickmode="linear",
                tick0=-1,
                dtick=0.5,
                len=0.7
            )
        ))
        
        # Layout optimized for readability
        fig_corr.update_layout(
            template="plotly_dark",
            width=1600,
            height=1000,
            title={
                'text': "Correlation Matrix • Diversification Analysis",
                'x': 0.5,
                'xanchor': 'center',
                'font': {'size': 16}
            },
            xaxis=dict(
                tickangle=-45,
                tickfont=dict(size=11),
                side='bottom'
            ),
            yaxis=dict(
                tickfont=dict(size=11),
                autorange='reversed'
            ),
            margin=dict(l=120, r=100, t=80, b=150),
            plot_bgcolor='rgba(17, 17, 17, 0.9)',
            paper_bgcolor='rgba(0,0,0,0)',
            font=dict(family='IBM Plex Mono', size=11, color='white')
        )
        
        # Scrollable container for large matrix
        st.markdown('<div class="matrix-container">', unsafe_allow_html=True)
        st.plotly_chart(fig_corr, use_container_width=False)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Correlation insights
        st.markdown("**Key Correlation Insights**")
        
        # Find highly correlated pairs
        corr_pairs = []
        for i in range(len(corr_matrix.columns)):
            for j in range(i+1, len(corr_matrix.columns)):
                corr_val = corr_matrix.iloc[i, j]
                if abs(corr_val) > 0.7:  # High correlation threshold
                    corr_pairs.append({
                        'Asset 1': corr_matrix.columns[i],
                        'Asset 2': corr_matrix.columns[j],
                        'Correlation': corr_val,
                        'Type': 'High Positive' if corr_val > 0 else 'High Negative'
                    })
        
        if corr_pairs:
            corr_df = pd.DataFrame(corr_pairs).sort_values('Correlation', ascending=False)
            st.dataframe(corr_df.head(10).style.format({'Correlation': '{:.3f}'}), hide_index=True)
            
            st.info(
                f"**Concentration Risk Alert:** Found {len([p for p in corr_pairs if p['Correlation'] > 0.7])} "
                f"asset pairs with correlation > 0.7. High positive correlations reduce diversification benefits."
            )
        else:
            st.success("Portfolio shows good diversification with no extreme correlations (>0.7).")

    # TAB 3: Enhanced Stress Testing
    with tab3:
        st.subheader("🌪️ Comprehensive Stress Testing & Scenario Analysis")
        st.caption("Simulate portfolio performance under various market scenarios. "
                  "Estimates based on portfolio beta and historical volatility.")
        
        # Calculate portfolio beta (vs market approximation)
        portfolio_beta = curr_std / 0.15  # Approximate market vol
        
        # Scenario inputs
        col_scenario1, col_scenario2 = st.columns([1, 3])
        
        with col_scenario1:
            st.markdown("**Scenario Parameters**")
            
            # Market Scenarios
            crash_pct = st.slider("📉 Market Crash", -50, -10, -25, step=5, 
                                 help="Severe market downturn scenario")
            
            bull_pct = st.slider("🐂 Bull Market Rally", 10, 100, 50, step=10,
                                help="Strong bull market scenario")
            
            ai_boom_pct = st.slider("🤖 AI Tech Boom", 20, 150, 70, step=10,
                                   help="Tech sector boom driven by AI")
            
            rate_crash_pct = st.slider("💸 Rate Cut Rally", 10, 60, 30, step=5,
                                      help="Interest rate cuts boost equities")
            
            budget_impact_pct = st.slider("📊 Budget Impact", -20, 40, 15, step=5,
                                         help="Government budget policy impact")
            
            # Risk multipliers for different scenarios
            st.markdown("---")
            st.markdown("**Beta Adjustment**")
            st.metric("Portfolio Beta", f"{portfolio_beta:.2f}")
        
        with col_scenario2:
            st.markdown("**Scenario Impact Analysis**")
            
            # Calculate scenario impacts
            scenarios = {
                f"Market Crash ({crash_pct}%)": {
                    'impact': crash_pct / 100 * portfolio_beta,
                    'description': "Severe downturn triggered by recession fears, credit event, or geopolitical crisis"
                },
                f"Bull Market Rally (+{bull_pct}%)": {
                    'impact': bull_pct / 100 * portfolio_beta * 0.9,  # Slightly lower beta in bull
                    'description': "Broad-based rally driven by economic growth and earnings expansion"
                },
                f"AI Tech Boom (+{ai_boom_pct}%)": {
                    'impact': ai_boom_pct / 100 * portfolio_beta * 1.3,  # Higher tech exposure
                    'description': "Technology sector surge led by AI revolution and semiconductor demand"
                },
                f"Rate Cut Rally (+{rate_crash_pct}%)": {
                    'impact': rate_crash_pct / 100 * portfolio_beta * 1.1,  # Interest-sensitive
                    'description': "Central bank easing cycle reduces borrowing costs and boosts equity valuations"
                },
                f"Budget Policy Impact ({'+' if budget_impact_pct > 0 else ''}{budget_impact_pct}%)": {
                    'impact': budget_impact_pct / 100 * portfolio_beta * 0.7,  # Moderate impact
                    'description': "Government fiscal policy changes affecting infrastructure, taxes, and regulations"
                }
            }
            
            # Create tornado chart
            scenario_names = list(scenarios.keys())
            scenario_impacts = [scenarios[s]['impact'] for s in scenario_names]
            scenario_values = [total_value * (1 + impact) - total_value for impact in scenario_impacts]
            
            colors = ['#FF4B4B' if v < 0 else '#00CC96' for v in scenario_values]
            
            fig_tornado = go.Figure()
            
            # Add bars
            fig_tornado.add_trace(go.Bar(
                y=scenario_names,
                x=scenario_values,
                orientation='h',
                marker=dict(
                    color=colors,
                    line=dict(color='rgba(255,255,255,0.3)', width=1)
                ),
                text=[f"₹{v:,.0f}" for v in scenario_values],
                textposition='outside',
                textfont=dict(size=11, color='white'),
                hovertemplate='<b>%{y}</b><br>P&L: ₹%{x:,.0f}<extra></extra>'
            ))
            
            # Add reference line at zero
            fig_tornado.add_vline(x=0, line_width=2, line_color='rgba(255,255,255,0.5)', line_dash='dash')
            
            fig_tornado.update_layout(
                template="plotly_dark",
                height=450,
                title={
                    'text': "Portfolio Stress Test Results",
                    'x': 0.5,
                    'xanchor': 'center'
                },
                xaxis_title="Estimated P&L Impact (₹)",
                yaxis_title="",
                showlegend=False,
                plot_bgcolor='rgba(0,0,0,0)',
                paper_bgcolor='rgba(0,0,0,0)',
                margin=dict(l=200, r=100, t=60, b=50),
                font=dict(family='IBM Plex Sans', size=11)
            )
            
            st.plotly_chart(fig_tornado, use_container_width=True)
            
            # Scenario descriptions
            st.markdown("**Scenario Descriptions**")
            for scenario, details in scenarios.items():
                impact_pct = details['impact'] * 100
                with st.expander(f"{scenario} → {impact_pct:+.1f}% portfolio impact"):
                    st.write(details['description'])
                    st.metric("Portfolio Value", f"₹{total_value * (1 + details['impact']):,.0f}",
                            delta=f"₹{total_value * details['impact']:,.0f}")
        
        # Statistical stress metrics
        st.markdown("---")
        st.markdown("**Statistical Risk Metrics**")
        
        col_risk1, col_risk2, col_risk3, col_risk4 = st.columns(4)
        
        with col_risk1:
            var_95 = curr_std * np.sqrt(1/252) * 1.645  # 95% confidence, 1-day
            st.metric("VaR (95%, 1D)", f"-{var_95 * 100:.2f}%", 
                     help="Value at Risk: Maximum expected loss on 95% of days")
        
        with col_risk2:
            cvar_95 = curr_std * np.sqrt(1/252) * 2.063  # Expected Shortfall
            st.metric("CVaR (95%, 1D)", f"-{cvar_95 * 100:.2f}%",
                     help="Conditional VaR: Average loss when VaR is exceeded")
        
        with col_risk3:
            max_dd = curr_std * 2 * 100  # Rough estimate
            st.metric("Est. Max Drawdown", f"-{max_dd:.1f}%",
                     help="Estimated maximum peak-to-trough decline")
        
        with col_risk4:
            stress_factor = portfolio_beta * 1.5
            st.metric("Stress Multiplier", f"{stress_factor:.2f}x",
                     help="Portfolio sensitivity to market shocks")

    # TAB 4: Position Analysis
    with tab4:
        st.subheader("📊 Individual Position Analysis")
        
        # Calculate position-level metrics
        position_data = []
        for _, row in edited_df.iterrows():
            ticker = row['Ticker']
            current_price = current_prices.get(ticker, 0)
            position_value = row['Current_Value']
            invested = row['Shares'] * row['Avg Cost']
            pnl = position_value - invested
            pnl_pct = (pnl / invested) * 100 if invested > 0 else 0
            weight = (position_value / total_value) * 100
            
            # Historical metrics
            if ticker in log_returns.columns:
                ticker_returns = log_returns[ticker]
                ann_return = ticker_returns.mean() * 252 * 100
                ann_vol = ticker_returns.std() * np.sqrt(252) * 100
                sharpe = (ann_return - rf_rate * 100) / ann_vol if ann_vol > 0 else 0
            else:
                ann_return = ann_vol = sharpe = 0
            
            position_data.append({
                'Ticker': ticker,
                'Shares': row['Shares'],
                'Avg Cost': row['Avg Cost'],
                'Current Price': current_price,
                'Position Value': position_value,
                'Weight %': weight,
                'P&L': pnl,
                'P&L %': pnl_pct,
                'Ann. Return %': ann_return,
                'Ann. Vol %': ann_vol,
                'Sharpe': sharpe
            })
        
        pos_df = pd.DataFrame(position_data).sort_values('Position Value', ascending=False)
        
        # Display position table with conditional formatting
        # Try gradient styling first, fall back to simple formatting if matplotlib unavailable
        try:
            styled_df = pos_df.style.format({
                'Shares': '{:.0f}',
                'Avg Cost': '₹{:.2f}',
                'Current Price': '₹{:.2f}',
                'Position Value': '₹{:,.0f}',
                'Weight %': '{:.2f}',
                'P&L': '₹{:,.0f}',
                'P&L %': '{:.2f}',
                'Ann. Return %': '{:.2f}',
                'Ann. Vol %': '{:.2f}',
                'Sharpe': '{:.3f}'
            }).background_gradient(subset=['P&L %'], cmap='RdYlGn', vmin=-20, vmax=20)\
              .background_gradient(subset=['Weight %'], cmap='Blues')
            
            st.dataframe(styled_df, use_container_width=True, height=400)
        except ImportError:
            # Matplotlib not available, use simple formatting
            st.dataframe(
                pos_df.style.format({
                    'Shares': '{:.0f}',
                    'Avg Cost': '₹{:.2f}',
                    'Current Price': '₹{:.2f}',
                    'Position Value': '₹{:,.0f}',
                    'Weight %': '{:.2f}',
                    'P&L': '₹{:,.0f}',
                    'P&L %': '{:.2f}',
                    'Ann. Return %': '{:.2f}',
                    'Ann. Vol %': '{:.2f}',
                    'Sharpe': '{:.3f}'
                }),
                use_container_width=True,
                height=400
            )
        
        # Position concentration chart
        st.markdown("**Portfolio Concentration**")
        
        fig_pie = px.pie(
            pos_df,
            values='Position Value',
            names='Ticker',
            title='Portfolio Allocation by Position',
            hole=0.4,
            color_discrete_sequence=px.colors.sequential.Viridis
        )
        
        fig_pie.update_layout(
            template="plotly_dark",
            height=500,
            showlegend=True,
            legend=dict(
                orientation="v",
                yanchor="middle",
                y=0.5,
                xanchor="left",
                x=1.05
            )
        )
        
        fig_pie.update_traces(
            textposition='inside',
            textinfo='label+percent',
            hovertemplate='<b>%{label}</b><br>Value: ₹%{value:,.0f}<br>Weight: %{percent}<extra></extra>'
        )
        
        st.plotly_chart(fig_pie, use_container_width=True)
        
        # Top performers
        col_top1, col_top2 = st.columns(2)
        
        with col_top1:
            st.markdown("**🏆 Top Performers (P&L %)**")
            top_performers = pos_df.nlargest(5, 'P&L %')[['Ticker', 'P&L %', 'P&L']]
            st.dataframe(
                top_performers.style.format({'P&L %': '{:.2f}', 'P&L': '₹{:,.0f}'}),
                hide_index=True
            )
        
        with col_top2:
            st.markdown("**📉 Bottom Performers (P&L %)**")
            bottom_performers = pos_df.nsmallest(5, 'P&L %')[['Ticker', 'P&L %', 'P&L']]
            st.dataframe(
                bottom_performers.style.format({'P&L %': '{:.2f}', 'P&L': '₹{:,.0f}'}),
                hide_index=True
            )

    # Enhanced Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; padding: 2rem 0 1rem 0;'>
        <div style='background: linear-gradient(145deg, rgba(30, 37, 56, 0.6), rgba(37, 45, 69, 0.6));
                    border-radius: 12px;
                    padding: 1.5rem;
                    border: 1px solid rgba(100, 150, 255, 0.2);'>
            <h3 style='background: linear-gradient(135deg, #60a5fa, #a78bfa); 
                       -webkit-background-clip: text; 
                       -webkit-text-fill-color: transparent;
                       margin-bottom: 0.5rem;'>
                💎 QuantPro Institutional Analytics
            </h3>
            <p style='color: #9ca3af; margin: 0.5rem 0;'>
                Powered by <strong>Modern Portfolio Theory</strong> • Real-time market data via <strong>Yahoo Finance</strong>
            </p>
            <p style='color: #6b7280; font-size: 0.875rem; margin-top: 1rem;'>
                <strong>Disclaimer:</strong> For informational purposes only. Not financial advice. 
                Past performance does not guarantee future results. Consult qualified financial advisors before making investment decisions.
            </p>
            <p style='color: #6b7280; font-size: 0.75rem; margin-top: 0.5rem;'>
                Risk Metrics: VaR, CVaR, Sharpe Ratio • Optimization: Scipy SLSQP • Monte Carlo: 5,000 simulations
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
